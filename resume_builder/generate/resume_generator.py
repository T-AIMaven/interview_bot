from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import json
from docx import Document
import pypandoc  # Ensure this is installed for PDF conversion
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_custom_margins(section, top, bottom, left, right):
    """Set custom margins for a section in inches."""
    section_properties = section._sectPr  # Access the section properties
    pgMar = section_properties.xpath('./w:pgMar')[0]  # Find the <w:pgMar> element

    # Convert inches to twips (1 inch = 1440 twips)
    pgMar.set(qn('w:top'), str(int(top * 1440)))
    pgMar.set(qn('w:bottom'), str(int(bottom * 1440)))
    pgMar.set(qn('w:left'), str(int(left * 1440)))
    pgMar.set(qn('w:right'), str(int(right * 1440)))


def create_resume_from_json(resume_json: dict, output_dir: str, output_format: str = 'docx'):
    if isinstance(resume_json, str):
        try:
            resume_json = json.loads(resume_json)
        except json.JSONDecodeError as e:
            raise ValueError(f"Invalid JSON format: {e}")

    doc = Document()

    # === Set Custom Margins ===
    section = doc.sections[0]
    set_custom_margins(section, top=1, bottom=1, left=0.8, right=0.8)


    # Define custom styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Georgia'  # Change the default font style
    font.size = Pt(10)   # Change the default font size

    # Name and Title
    name_para = doc.add_paragraph()
    
    

    name_run = name_para.add_run(resume_json['name'])
    name_run.bold = True
    name_run.font.size = Pt(22)
    name_run.font.name = 'Georgia'  # Custom font for the name
    name_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    title_para = doc.add_paragraph()
    title_run = title_para.add_run(resume_json['title'])
    title_run.font.size = Pt(12)
    title_run.font.name = 'Georgia'  # Custom font for the title
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Contact Info
    contact_info = resume_json['Contact']
    contact_para = doc.add_paragraph()
    if 'email' in contact_info:
        email_run = contact_para.add_run("Email: ")
        email_run.bold = True
        email_run.font.name = 'Georgia'
        contact_para.add_run(contact_info['email'])

    if 'phone' in contact_info:
        contact_para.add_run(", ")  # Add a new line
        phone_run = contact_para.add_run("Phone: ")
        phone_run.bold = True
        phone_run.font.name = 'Georgia'
        contact_para.add_run(contact_info['phone'])

    if 'location' in contact_info:
        contact_para.add_run(", ")
        location_run = contact_para.add_run("Location: ")
        location_run.bold = True
        location_run.font.name = 'Georgia'
        contact_para.add_run(contact_info['location'])

    if 'linkedin' in contact_info:
        contact_para.add_run(", ")
        linkedin_run = contact_para.add_run("LinkedIn: ")
        linkedin_run.bold = True
        linkedin_run.font.name = 'Georgia'
        contact_para.add_run(contact_info['linkedin'])

    if 'website' in contact_info:
        contact_para.add_run(", ")
        website_run = contact_para.add_run("Website: ")
        website_run.bold = True
        website_run.font.name = 'Georgia'
        contact_para.add_run(contact_info['website'])

    contact_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add a line break
    # doc.add_paragraph()

    # Summary Section
    summary_heading = doc.add_paragraph('SUMMARY', style='Heading 2')
    summary_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    summary_heading.runs[0].font.name = 'Georgia'  # Change font style for the heading
    summary_heading.runs[0].font.size = Pt(14)

    # Add a paragraph for the summary text
    summary = doc.add_paragraph()
    summary_run = summary.add_run(resume_json['summary'])  # Add a Run to the paragraph
    summary_run.font.name = 'Georgia'  # Change font style for the summary text
    summary_run.font.size = Pt(11)
    summary.paragraph_format.space_after = Pt(10)

    # Add a line break
    # doc.add_paragraph()

    

    # Skills Section
    skills_heading = doc.add_paragraph('SKILLS', style='Heading 2')
    skills_heading.runs[0].font.name = 'Georgia'
    skills_heading.runs[0].font.size = Pt(14)

    for skill, items in resume_json['skills'].items():
        para = doc.add_paragraph()
        skill_run = para.add_run(f"{skill}: ")
        skill_run.bold = True
        skill_run.font.name = 'Georgia'
        para.add_run(", ".join(items))

    # Add a line break
    doc.add_paragraph()
    
    # Experience Section
    experience_heading = doc.add_paragraph('PROFESSIONAL EXPERIENCE', style='Heading 2')
    experience_heading.runs[0].font.name = 'Georgia'
    experience_heading.runs[0].font.size = Pt(14)

    for exp in resume_json['experiences']:
        heading = doc.add_paragraph(f"{exp['company']} - {exp['role']}", style='Heading 3')
        heading.paragraph_format.space_after = Pt(1)
        heading.runs[0].font.name = 'Georgia'
        heading.runs[0].font.size = Pt(12)

        date_para = doc.add_paragraph(f"{exp['start_date']} - {exp['end_date']}")
        date_para.paragraph_format.space_after = Pt(3)
        date_para.runs[0].font.name = 'Georgia'
        date_para.runs[0].font.size = Pt(11)

        for bullet in exp['bullets']:
            bullet_para = doc.add_paragraph(bullet, style='List Bullet')
            bullet_para.paragraph_format.space_after = Pt(1)
            bullet_para.runs[0].font.name = 'Georgia'
            bullet_para.runs[0].font.size = Pt(11)

    # Add a line break
    # doc.add_paragraph()
    
    # Education Section
    education_heading = doc.add_paragraph('EDUCATION', style='Heading 2')
    education_heading.runs[0].font.name = 'Georgia'
    education_heading.runs[0].font.size = Pt(14)

    for edu in resume_json['education']:
        edu_heading = doc.add_paragraph(f"{edu['institute_name']} - {edu['degree']}", style='Heading 3')
        edu_heading.paragraph_format.space_after = Pt(1)
        edu_heading.runs[0].font.name = 'Georgia'
        edu_heading.runs[0].font.size = Pt(12)

        edu_dates = doc.add_paragraph(f"{edu['start_date']} - {edu['end_date']}")
        edu_dates.paragraph_format.space_after = Pt(1)
        edu_dates.runs[0].font.name = 'Georgia'
        edu_dates.runs[0].font.size = Pt(11)

        if 'gpa' in edu:
            gpa = doc.add_paragraph(f"GPA: {edu['gpa']}")
            gpa.paragraph_format.space_after = Pt(5)
            gpa.runs[0].font.name = 'Georgia'
            gpa.runs[0].font.size = Pt(11)

    # Ensure the output directory exists
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # Save the DOCX
    doc_filename = f"{resume_json['name']} Resume.docx"
    doc_path = os.path.join(output_dir, doc_filename)
    doc.save(doc_path)

    # Convert to PDF
    pdf_path = None
    # if output_format == 'pdf' or output_format == 'both':
    #     pdf_filename = f"{resume_json['name']} Resume.pdf"
    #     pdf_path = os.path.join(output_dir, pdf_filename)
    #     try:
    #         pypandoc.convert_file(doc_path, 'pdf', outputfile=pdf_path)
    #     except Exception as e:
    #         raise IOError(f"Failed to convert DOCX to PDF: {e}")

    return doc_path, pdf_path